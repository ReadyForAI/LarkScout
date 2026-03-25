#!/usr/bin/env python3
# /// script
# requires-python = ">=3.10"
# dependencies = ["pymupdf", "python-docx", "google-genai", "Pillow", "fastapi", "uvicorn", "python-multipart"]
# ///

import asyncio
import hashlib
import json
import logging
import os
import re
import shutil
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from pydantic import BaseModel

from i18n import init_locale, prompt, t, tmpl

init_locale()

logger = logging.getLogger("larkscout_docreader")
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(name)s] %(message)s")


# ═══════════════════════════════════════════
# Data structures
# ═══════════════════════════════════════════

@dataclass
class PageContent:
    """Single page content."""
    page_num: int
    text: str
    is_ocr: bool = False
    tables: list[str] = field(default_factory=list)


@dataclass
class Section:
    """Document section."""
    index: int
    title: str
    level: int  # heading level 1-3
    text: str
    page_range: str  # "p.5-12"
    summary: str = ""
    sid: str = ""  # stable ID


@dataclass
class ParsedDocument:
    """Parsed document result."""
    filename: str
    file_type: str  # "pdf" | "docx"
    total_pages: int
    pages: list[PageContent]
    sections: list[Section]
    ocr_page_count: int = 0
    table_count: int = 0


# ═══════════════════════════════════════════
# LLM provider wrapper
# ═══════════════════════════════════════════


def gemini_ocr(image_bytes: bytes, page_num: int) -> str:
    """OCR a single page image via the active LLM provider."""
    from providers import get_provider

    return get_provider().ocr(image_bytes, page_num)


def gemini_summarize(text: str, summarize_prompt: str, max_retries: int = 2) -> str:
    """Generate summary via the active LLM provider."""
    from providers import get_provider

    return get_provider().summarize(text, summarize_prompt, max_retries=max_retries)


# ═══════════════════════════════════════════
# Token estimation
# ═══════════════════════════════════════════

def _estimate_tokens(text: str) -> int:
    """Rough token estimate. CJK ~2.5 chars/tok, Latin ~4 chars/tok."""
    if not text:
        return 0
    cjk_count = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    ratio = cjk_count / max(len(text), 1)
    chars_per_token = 2.5 * ratio + 4.0 * (1 - ratio)
    return int(len(text) / chars_per_token)


# ═══════════════════════════════════════════
# Smart OCR detection
# ═══════════════════════════════════════════

OCR_THRESHOLD = 50


def _parse_page_range(spec: str, total_pages: int) -> set[int]:
    """Parse page range spec: "10-30" or "5,10-15,20"."""
    pages = set()
    for part in spec.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            start = max(1, int(a.strip()))
            end = min(total_pages, int(b.strip()))
            pages.update(range(start, end + 1))
        else:
            p = int(part.strip())
            if 1 <= p <= total_pages:
                pages.add(p)
    return pages


def _should_ocr(page, text: str, threshold: int) -> bool:
    """
    Multi-signal OCR detection:
      Signal 1: too little text
      Signal 2: page has images and text is sparse (scan indicator)
      Signal 3: low useful-character ratio (garbled or mostly whitespace)
    """
    if len(text) < threshold:
        return True
    try:
        images = page.get_images(full=False)
        if len(images) > 0 and len(text) < threshold * 3:
            return True
    except Exception:
        pass
    if len(text) > 0:
        useful = sum(1 for c in text if c.isalnum() or '\u4e00' <= c <= '\u9fff')
        if useful / len(text) < 0.3 and len(text) < threshold * 5:
            return True
    return False


def _ocr_cache_path(doc_dir: Path, page_num: int) -> Path:
    cache_dir = doc_dir / ".cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    return cache_dir / f"ocr_p{page_num:04d}.txt"


def _ocr_cache_key(image_bytes: bytes) -> str:
    return hashlib.sha1(image_bytes).hexdigest()[:16]


# ═══════════════════════════════════════════
# Section stable ID
# ═══════════════════════════════════════════

def _section_sid(title: str, text: str) -> str:
    raw = (title + text[:200]).encode("utf-8")
    return hashlib.sha1(raw).hexdigest()[:12]


# ═══════════════════════════════════════════
# PDF parsing
# ═══════════════════════════════════════════

def parse_pdf(
    filepath: Path,
    force_ocr: bool = False,
    ocr_threshold: int = OCR_THRESHOLD,
    ocr_pages_spec: str | None = None,
    extract_tables: bool = True,
    max_tables_per_page: int = 3,
    concurrency: int = 3,
    cache_dir: Path | None = None,
) -> ParsedDocument:
    import fitz

    logger.info(f"Parsing PDF: {filepath.name}")
    doc = fitz.open(str(filepath))
    total_pages = len(doc)
    logger.info(f"Total pages: {total_pages}")

    ocr_page_set: set[int] | None = None
    if ocr_pages_spec:
        ocr_page_set = _parse_page_range(ocr_pages_spec, total_pages)
        logger.info(f"OCR target pages: {sorted(ocr_page_set)}")

    # PDF TOC
    toc = doc.get_toc(simple=True)
    if toc:
        logger.info(f"PDF TOC detected: {len(toc)} entries")

    pages: list[PageContent] = []
    ocr_count = 0
    table_count = 0
    table_hashes: set[str] = set()

    ocr_tasks: list[tuple[int, bytes]] = []
    ocr_results: dict[int, str] = {}

    for i, page in enumerate(doc):
        page_num = i + 1
        text = page.get_text("text").strip()
        tables: list[str] = []

        # Tables: optional + per-page limit + hash dedup
        if extract_tables:
            try:
                page_tables = page.find_tables()
                added = 0
                for tab in page_tables:
                    if added >= max_tables_per_page:
                        break
                    md_table = _table_to_markdown(tab.extract())
                    if md_table:
                        h = hashlib.md5(md_table.encode()).hexdigest()
                        if h not in table_hashes:
                            table_hashes.add(h)
                            tables.append(md_table)
                            table_count += 1
                            added += 1
            except Exception:
                pass

        # OCR decision
        if ocr_page_set is not None:
            need_ocr = page_num in ocr_page_set
        elif force_ocr:
            need_ocr = True
        else:
            need_ocr = _should_ocr(page, text, ocr_threshold)

        if need_ocr:
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")

            # OCR cache
            if cache_dir:
                cp = _ocr_cache_path(cache_dir, page_num)
                ck = _ocr_cache_key(img_bytes)
                ck_path = cp.with_suffix(f".{ck}.txt")
                if ck_path.exists():
                    text = ck_path.read_text(encoding="utf-8")
                    logger.info(f"Page {page_num}/{total_pages}: OCR cache hit")
                    ocr_count += 1
                    pages.append(PageContent(page_num=page_num, text=text, is_ocr=True, tables=tables))
                    continue

            ocr_tasks.append((page_num, img_bytes))
            pages.append(PageContent(page_num=page_num, text="", is_ocr=True, tables=tables))
            ocr_count += 1
        else:
            if page_num % 20 == 0 or page_num == total_pages:
                logger.info(f"Page {page_num}/{total_pages}: text extracted")
            pages.append(PageContent(page_num=page_num, text=text, is_ocr=False, tables=tables))

    doc.close()

    # Concurrent OCR
    if ocr_tasks:
        logger.info(f"Concurrent OCR: {len(ocr_tasks)} pages ({concurrency} workers)...")

        def _do_ocr(args):
            pn, img_b = args
            result = gemini_ocr(img_b, pn)
            return pn, img_b, result

        with ThreadPoolExecutor(max_workers=concurrency) as pool:
            futures = {pool.submit(_do_ocr, task): task for task in ocr_tasks}
            for fut in as_completed(futures):
                pn, img_b, result = fut.result()
                ocr_results[pn] = result
                logger.info(f"Page {pn}/{total_pages}: OCR done")
                if cache_dir:
                    cp = _ocr_cache_path(cache_dir, pn)
                    ck = _ocr_cache_key(img_b)
                    ck_path = cp.with_suffix(f".{ck}.txt")
                    ck_path.write_text(result, encoding="utf-8")

    for page in pages:
        if page.is_ocr and not page.text and page.page_num in ocr_results:
            page.text = ocr_results[page.page_num]

    # TOC-first split
    if toc:
        sections = _split_sections_from_toc(pages, toc)
    else:
        sections = _split_sections(pages)

    for sec in sections:
        sec.sid = _section_sid(sec.title, sec.text)

    logger.info(f"Parse complete: {len(sections)} sections, {ocr_count} OCR pages, {table_count} tables")

    return ParsedDocument(
        filename=filepath.name, file_type="pdf", total_pages=total_pages,
        pages=pages, sections=sections, ocr_page_count=ocr_count, table_count=table_count,
    )


# ═══════════════════════════════════════════
# Word parsing
# ═══════════════════════════════════════════

def parse_word(filepath: Path, extract_tables: bool = True) -> ParsedDocument:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    logger.info(f"Parsing Word: {filepath.name}")
    try:
        doc = Document(str(filepath))
    except PackageNotFoundError:
        raise RuntimeError(t("file_open_failed", path=str(filepath)))

    elements: list[dict] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        level = 0
        if style.startswith("Heading"):
            try:
                level = int(style.replace("Heading", "").strip())
            except ValueError:
                level = 1
        elements.append({"text": text, "level": level, "type": "paragraph"})

    table_count = 0
    if extract_tables:
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            md_table = _rows_to_markdown(rows)
            if md_table:
                elements.append({"text": md_table, "level": 0, "type": "table"})
                table_count += 1

    full_text = "\n\n".join(e["text"] for e in elements)
    est_pages = max(1, len(full_text) // 3000)
    pages = [PageContent(page_num=1, text=full_text)]
    sections = _split_sections_from_elements(elements)
    for sec in sections:
        sec.sid = _section_sid(sec.title, sec.text)

    logger.info(f"Parse complete: {len(sections)} sections, ~{est_pages} pages, {table_count} tables")
    return ParsedDocument(
        filename=filepath.name, file_type="docx", total_pages=est_pages,
        pages=pages, sections=sections, table_count=table_count,
    )


# ═══════════════════════════════════════════
# XLSX parsing
# ═══════════════════════════════════════════

def parse_xlsx(filepath: Path) -> ParsedDocument:
    """Parse an XLSX workbook; each sheet becomes one section and one Markdown table."""
    import openpyxl

    logger.info(f"Parsing XLSX: {filepath.name}")
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)

    sections: list[Section] = []
    pages: list[PageContent] = []
    table_count = 0

    for idx, sheet_name in enumerate(wb.sheetnames, 1):
        ws = wb[sheet_name]
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            str_row = [str(cell) if cell is not None else "" for cell in row]
            if any(c.strip() for c in str_row):
                rows.append(str_row)

        if not rows:
            continue

        md_table = _rows_to_markdown(rows)
        page = PageContent(page_num=idx, text=md_table, tables=[md_table] if md_table else [])
        pages.append(page)

        if md_table:
            table_count += 1

        sid = _section_sid(sheet_name, md_table)
        sections.append(Section(
            index=idx,
            title=sheet_name,
            level=1,
            text=md_table,
            page_range=f"sheet {idx}",
            sid=sid,
        ))

    wb.close()

    logger.info(f"XLSX parse complete: {len(sections)} sheets, {table_count} tables")
    return ParsedDocument(
        filename=filepath.name,
        file_type="xlsx",
        total_pages=max(len(pages), 1),
        pages=pages,
        sections=sections,
        table_count=table_count,
    )


# ═══════════════════════════════════════════
# CSV parsing
# ═══════════════════════════════════════════

def parse_csv(filepath: Path) -> ParsedDocument:
    """Parse a CSV file; the entire file becomes one section and one Markdown table."""
    import csv

    logger.info(f"Parsing CSV: {filepath.name}")

    # Try UTF-8-with-BOM first (common Excel export), fall back to latin-1
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            with open(filepath, newline="", encoding=encoding) as f:
                reader = csv.reader(f)
                rows: list[list[str]] = [
                    row for row in reader if any(cell.strip() for cell in row)
                ]
            break
        except UnicodeDecodeError:
            continue
    else:
        rows = []

    md_table = _rows_to_markdown(rows) if rows else ""
    page = PageContent(page_num=1, text=md_table, tables=[md_table] if md_table else [])

    stem = filepath.stem
    sid = _section_sid(stem, md_table)
    section = Section(
        index=1,
        title=stem,
        level=1,
        text=md_table,
        page_range="sheet 1",
        sid=sid,
    )

    table_count = 1 if md_table else 0
    logger.info(f"CSV parse complete: {len(rows)} rows, {table_count} tables")
    return ParsedDocument(
        filename=filepath.name,
        file_type="csv",
        total_pages=1,
        pages=[page],
        sections=[section] if md_table else [],
        table_count=table_count,
    )


# ═══════════════════════════════════════════
# Section splitting
# ═══════════════════════════════════════════

HEADING_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十\d]+[章节部分篇]\s*[、:：]?\s*.+"),
    re.compile(r"^[（(]?[一二三四五六七八九十]+[）)]?[、.．]\s*.+"),
    re.compile(r"^\d+(\.\d+)*[.、．)\s]\s*.{2,}"),
    re.compile(r"^[A-Z][A-Z\s]{5,}$"),
    re.compile(r"^(摘要|目录|引言|绪论|前言|导论|背景|概述|总结|结论|致谢|参考文献|附录|附件)$"),
]


def _is_heading(text: str) -> int:
    text = text.strip()
    if not text or len(text) > 100:
        return 0
    for i, pattern in enumerate(HEADING_PATTERNS):
        if pattern.match(text):
            return 1 if i < 2 else 2
    return 0


def _split_sections_from_toc(pages: list[PageContent], toc: list) -> list[Section]:
    """Split sections using PDF TOC."""
    if not toc or not pages:
        return _split_sections(pages)

    page_texts: dict[int, str] = {}
    for p in pages:
        t = p.text
        if p.tables:
            t += "\n\n" + "\n\n".join(p.tables)
        page_texts[p.page_num] = t

    max_page = max(p.page_num for p in pages)
    sections: list[Section] = []

    for i, (level, title, start_page) in enumerate(toc):
        end_page = toc[i + 1][2] - 1 if i + 1 < len(toc) else max_page
        end_page = max(end_page, start_page)
        text_parts = [page_texts[pn] for pn in range(start_page, end_page + 1) if pn in page_texts]
        text = "\n\n".join(text_parts).strip()
        if not text:
            continue
        sections.append(Section(
            index=len(sections) + 1, title=title.strip(),
            level=min(level, 3), text=text, page_range=f"p.{start_page}-{end_page}",
        ))

    if len(sections) < 2:
        logger.warning("PDF TOC produced too few sections, falling back to regex split")
        return _split_sections(pages)
    return sections


def _split_sections(pages: list[PageContent]) -> list[Section]:
    sections: list[Section] = []
    current_title = tmpl("default_section_title")
    current_level = 1
    current_lines: list[str] = []
    current_start_page = 1
    sec_index = 0

    for page in pages:
        for line in page.text.split("\n"):
            line = line.strip()
            if not line:
                continue
            heading_level = _is_heading(line)
            if heading_level > 0 and current_lines:
                sec_index += 1
                sections.append(Section(
                    index=sec_index, title=current_title, level=current_level,
                    text="\n".join(current_lines),
                    page_range=f"p.{current_start_page}-{page.page_num}",
                ))
                current_title = line
                current_level = heading_level
                current_lines = []
                current_start_page = page.page_num
            else:
                current_lines.append(line)
        for table in page.tables:
            current_lines.append(f"\n{table}\n")

    if current_lines:
        sec_index += 1
        last_page = pages[-1].page_num if pages else 1
        sections.append(Section(
            index=sec_index, title=current_title, level=current_level,
            text="\n".join(current_lines), page_range=f"p.{current_start_page}-{last_page}",
        ))

    if not sections:
        full_text = "\n\n".join(p.text for p in pages)
        sections.append(Section(
            index=1, title=tmpl("full_document_title"), level=1, text=full_text,
            page_range=f"p.1-{pages[-1].page_num if pages else 1}",
        ))
    return sections


def _split_sections_from_elements(elements: list[dict]) -> list[Section]:
    sections: list[Section] = []
    current_title = tmpl("default_section_title")
    current_level = 1
    current_lines: list[str] = []
    sec_index = 0

    for elem in elements:
        if elem["level"] > 0:
            if current_lines:
                sec_index += 1
                sections.append(Section(
                    index=sec_index, title=current_title,
                    level=current_level, text="\n".join(current_lines), page_range="",
                ))
            current_title = elem["text"]
            current_level = elem["level"]
            current_lines = []
        else:
            heading_level = _is_heading(elem["text"]) if elem["type"] == "paragraph" else 0
            if heading_level > 0 and current_lines:
                sec_index += 1
                sections.append(Section(
                    index=sec_index, title=current_title,
                    level=current_level, text="\n".join(current_lines), page_range="",
                ))
                current_title = elem["text"]
                current_level = heading_level
                current_lines = []
            else:
                current_lines.append(elem["text"])

    if current_lines:
        sec_index += 1
        sections.append(Section(
            index=sec_index, title=current_title,
            level=current_level, text="\n".join(current_lines), page_range="",
        ))
    if not sections:
        full_text = "\n".join(e["text"] for e in elements)
        sections.append(Section(index=1, title=tmpl("full_document_title"), level=1, text=full_text, page_range=""))
    return sections


# ═══════════════════════════════════════════
# Table utilities
# ═══════════════════════════════════════════

def _table_to_markdown(table_data: list[list]) -> str:
    if not table_data or len(table_data) < 2:
        return ""
    return _rows_to_markdown([[cell if cell else "" for cell in row] for row in table_data])


def _rows_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    clean_rows = [[cell.replace("\n", " ").strip() for cell in row] for row in rows]
    if not clean_rows:
        return ""
    header = "| " + " | ".join(clean_rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * len(clean_rows[0])) + " |"
    body_lines = []
    for row in clean_rows[1:]:
        while len(row) < len(clean_rows[0]):
            row.append("")
        body_lines.append("| " + " | ".join(row[:len(clean_rows[0])]) + " |")
    return "\n".join([header, sep] + body_lines)


# ═══════════════════════════════════════════
# Summary generation
# ═══════════════════════════════════════════

SUMMARY_MAX_CHARS = 500


def generate_summaries(parsed: ParsedDocument, concurrency: int = 3) -> tuple[str, str, list[Section]]:
    logger.info("Generating summaries...")

    # Dynamic batching by token estimate
    BATCH_TOKEN_LIMIT = 10000
    batches: list[list[Section]] = []
    current_batch: list[Section] = []
    current_tokens = 0

    for sec in parsed.sections:
        sec_tokens = _estimate_tokens(sec.text) + _estimate_tokens(sec.title) + 20
        if current_tokens + sec_tokens > BATCH_TOKEN_LIMIT and current_batch:
            batches.append(current_batch)
            current_batch = []
            current_tokens = 0
        current_batch.append(sec)
        current_tokens += sec_tokens
    if current_batch:
        batches.append(current_batch)

    # Concurrent summarization
    if len(batches) > 1 and concurrency > 1:
        logger.info(f"{len(batches)} batches, {min(concurrency, len(batches))} workers")
        with ThreadPoolExecutor(max_workers=min(concurrency, len(batches))) as pool:
            futures = {pool.submit(_summarize_batch, batch): batch for batch in batches}
            for fut in as_completed(futures):
                fut.result()
    else:
        for batch in batches:
            _summarize_batch(batch)

    logger.info(f"{len(parsed.sections)} section summaries complete")

    # Brief input compression
    if len(parsed.sections) > 60:
        sections_overview = _compress_sections_for_brief(parsed.sections)
    else:
        sections_overview = "\n\n".join(
            f"## {sec.title} ({sec.page_range})\n{sec.summary[:SUMMARY_MAX_CHARS]}"
            for sec in parsed.sections if sec.summary
        )

    brief = gemini_summarize(
        f"Document: {parsed.filename}\nTotal pages: {parsed.total_pages}\n\n{sections_overview}",
        prompt("brief"),
    )
    logger.info("Brief generation complete")

    digest = gemini_summarize(
        f"Document: {parsed.filename}\n\nBriefing:\n{brief}",
        prompt("digest"),
    )
    logger.info("Digest generation complete")

    return digest, brief, parsed.sections


def _summarize_batch(sections: list[Section]):
    """Batch summarize with JSON output + single fallback."""
    n = len(sections)

    if n == 1:
        sec = sections[0]
        sec.summary = gemini_summarize(
            f"## {sec.title} ({sec.page_range})\n\n{sec.text}",
            prompt("section_summary"),
        )
        logger.info(f"Section {sec.index}: {sec.title[:30]}... done")
        return

    batch_text = ""
    for sec in sections:
        batch_text += f"\n\n## Section {sec.index}: {sec.title} ({sec.page_range})\n\n{sec.text}"

    result = gemini_summarize(batch_text, prompt("batch_summary", n=n))

    # JSON parse
    parsed_ok = False
    try:
        clean = result.strip()
        if clean.startswith("```"):
            clean = re.sub(r"^```(?:json)?\s*", "", clean)
            clean = re.sub(r"\s*```$", "", clean)
        items = json.loads(clean)
        if isinstance(items, list) and len(items) >= n:
            for sec in sections:
                match = next((it for it in items if it.get("index") == sec.index), None)
                if match and match.get("summary"):
                    sec.summary = match["summary"]
                else:
                    sec.summary = t("summary_missing")
            parsed_ok = True
    except (json.JSONDecodeError, KeyError, TypeError):
        pass

    if parsed_ok:
        for sec in sections:
            logger.info(f"Section {sec.index}: {sec.title[:30]}... done")
        return

    # Fallback
    logger.warning(f"Batch JSON parse failed, falling back to single ({n} items)")
    for sec in sections:
        sec.summary = gemini_summarize(
            f"## {sec.title} ({sec.page_range})\n\n{sec.text}",
            prompt("section_summary"),
        )
        logger.info(f"Section {sec.index}: {sec.title[:30]}... done (single)")


def _compress_sections_for_brief(sections: list[Section]) -> str:
    groups = []
    for i in range(0, len(sections), 10):
        group = sections[i:i + 10]
        group_text = "; ".join(f"{s.title}: {s.summary[:150]}" for s in group if s.summary)
        groups.append(f"**Sections {group[0].index}-{group[-1].index}**: {group_text}")
    return "\n\n".join(groups)


# ═══════════════════════════════════════════
# Output file writing
# ═══════════════════════════════════════════

def write_output(doc_id: str, parsed: ParsedDocument, digest: str, brief: str, output_dir: Path,
                 tags: list[str] | None = None, source: str = "upload",
                 original_path: str | None = None):
    doc_dir = output_dir / doc_id
    sections_dir = doc_dir / "sections"
    tables_dir = doc_dir / "tables"
    doc_dir.mkdir(parents=True, exist_ok=True)
    sections_dir.mkdir(exist_ok=True)

    meta = {
        "doc_id": doc_id,
        "filename": parsed.filename,
        "file_type": parsed.file_type,
        "total_pages": parsed.total_pages,
        "section_count": len(parsed.sections),
        "ocr_page_count": parsed.ocr_page_count,
        "table_count": parsed.table_count,
        "created_at": datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "sections": [
            {"index": sec.index, "sid": sec.sid, "title": sec.title,
             "page_range": sec.page_range, "char_count": len(sec.text)}
            for sec in parsed.sections
        ],
    }
    _write_json(doc_dir / ".meta.json", meta)
    logger.info(".meta.json written")

    _write_text(doc_dir / "digest.md", f"# {doc_id}: {parsed.filename}\n\n{digest}\n")
    logger.info("digest.md written")

    brief_header = tmpl("brief_header",
        doc_id=doc_id, filename=parsed.filename, pages=parsed.total_pages,
        sections=len(parsed.sections), ocr_pages=parsed.ocr_page_count)
    _write_text(doc_dir / "brief.md", brief_header + brief + "\n")
    logger.info("brief.md written")

    full_parts = [f"{'#' * min(sec.level + 1, 4)} {sec.title}\n\n{sec.text}" for sec in parsed.sections]
    _write_text(doc_dir / "full.md", f"# {parsed.filename}\n\n" + "\n\n---\n\n".join(full_parts) + "\n")
    logger.info("full.md written")

    for sec in parsed.sections:
        sec_filename = f"{sec.index:02d}-{sec.sid}-{_safe_filename(sec.title)}.md"
        sec_content = tmpl("section_header", title=sec.title, index=sec.index, sid=sec.sid, page_range=sec.page_range)
        if sec.summary:
            sec_content += tmpl("section_summary_line", summary=sec.summary)
        sec_content += sec.text + "\n"
        _write_text(sections_dir / sec_filename, sec_content)
    logger.info(f"sections/ ({len(parsed.sections)} files)")

    all_tables = [(p.page_num, t) for p in parsed.pages for t in p.tables]
    if all_tables:
        tables_dir.mkdir(exist_ok=True)
        for i, (page_num, table_md) in enumerate(all_tables, 1):
            _write_text(tables_dir / f"table-{i:02d}.md", f"# Table {i} (page {page_num})\n\n{table_md}\n")
        logger.info(f"tables/ ({len(all_tables)} files)")

    # v3: content_hash
    full_text = "\n".join(sec.text for sec in parsed.sections)
    content_hash = "sha256:" + hashlib.sha256(full_text.encode("utf-8", errors="ignore")).hexdigest()

    # manifest.json + v3 provenance
    manifest = {
        "doc_id": doc_id,
        "filename": parsed.filename,
        "file_type": parsed.file_type,
        "source": source,
        "paths": {"digest": "digest.md", "brief": "brief.md", "full": "full.md", "sections_dir": "sections/"},
        "sections": [
            {
                "sid": sec.sid, "index": sec.index, "title": sec.title,
                "page_range": sec.page_range, "char_count": len(sec.text),
                "summary_preview": (sec.summary[:120] + "...") if len(sec.summary) > 120 else sec.summary,
                "file": f"sections/{sec.index:02d}-{sec.sid}-{_safe_filename(sec.title)}.md",
            }
            for sec in parsed.sections
        ],
        "provenance": {
            "source": source,
            "original_path": original_path or str(parsed.filename),
            "upload_time": meta["created_at"],
            "content_hash": content_hash,
        },
    }
    _write_json(doc_dir / "manifest.json", manifest)
    logger.info("manifest.json written")

    _update_doc_index(output_dir, meta, digest, tags=tags, source=source, content_hash=content_hash)


def write_output_extract_only(doc_id: str, parsed: ParsedDocument, output_dir: Path,
                              tags: list[str] | None = None, source: str = "upload"):
    doc_dir = output_dir / doc_id
    sections_dir = doc_dir / "sections"
    doc_dir.mkdir(parents=True, exist_ok=True)
    sections_dir.mkdir(exist_ok=True)

    meta = {
        "doc_id": doc_id, "filename": parsed.filename, "file_type": parsed.file_type,
        "total_pages": parsed.total_pages, "section_count": len(parsed.sections),
        "ocr_page_count": parsed.ocr_page_count, "table_count": parsed.table_count,
        "created_at": datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "sections": [
            {"index": sec.index, "sid": sec.sid, "title": sec.title,
             "page_range": sec.page_range, "char_count": len(sec.text)}
            for sec in parsed.sections
        ],
    }
    _write_json(doc_dir / ".meta.json", meta)

    full_parts = [f"{'#' * min(sec.level + 1, 4)} {sec.title}\n\n{sec.text}" for sec in parsed.sections]
    _write_text(doc_dir / "full.md", f"# {parsed.filename}\n\n" + "\n\n---\n\n".join(full_parts) + "\n")

    for sec in parsed.sections:
        fn = f"{sec.index:02d}-{sec.sid}-{_safe_filename(sec.title)}.md"
        _write_text(sections_dir / fn, f"# {sec.title}\n\n{sec.text}\n")

    _write_text(doc_dir / "digest.md", f"{tmpl('digest_title', doc_id=doc_id, filename=parsed.filename)}\n\n{t('summary_pending')}\n")
    _write_text(doc_dir / "brief.md", f"{tmpl('digest_title', doc_id=doc_id, filename=parsed.filename)}\n\n{t('summary_pending')}\n")

    manifest = {
        "doc_id": doc_id, "filename": parsed.filename,
        "paths": {"digest": "digest.md", "brief": "brief.md", "full": "full.md", "sections_dir": "sections/"},
        "sections": [
            {"sid": sec.sid, "index": sec.index, "title": sec.title,
             "page_range": sec.page_range, "char_count": len(sec.text),
             "summary_preview": "",
             "file": f"sections/{sec.index:02d}-{sec.sid}-{_safe_filename(sec.title)}.md"}
            for sec in parsed.sections
        ],
    }
    _write_json(doc_dir / "manifest.json", manifest)
    _update_doc_index(output_dir, meta, t("summary_pending"), tags=tags, source=source)
    logger.info(f"Text extraction complete (no summary): {doc_dir}")


# ═══════════════════════════════════════════
# Document index
# ═══════════════════════════════════════════

def _update_doc_index(docs_dir: Path, meta: dict, digest: str,
                      tags: list[str] | None = None,
                      source: str = "upload",
                      source_url: str | None = None,
                      content_hash: str | None = None):
    """Update doc-index.json (v2 format with source/tags/content_hash)."""
    index_path = docs_dir / "doc-index.json"
    if index_path.exists():
        try:
            with open(index_path, encoding="utf-8") as f:
                index = json.load(f)
        except (json.JSONDecodeError, Exception):
            index = {"version": 2, "documents": []}
    else:
        index = {"version": 2, "documents": []}

    index["version"] = 2
    index["documents"] = [d for d in index["documents"] if d.get("id") != meta["doc_id"]]

    entry = {
        "id": meta["doc_id"], "filename": meta["filename"], "file_type": meta["file_type"],
        "source": source,
        "pages": meta["total_pages"], "sections": meta["section_count"],
        "ocr_pages": meta.get("ocr_page_count", 0), "tables": meta.get("table_count", 0),
        "digest": digest[:200], "digest_path": f"docs/{meta['doc_id']}/digest.md",
        "tags": tags or [],
        "created_at": meta["created_at"],
    }
    if content_hash:
        entry["content_hash"] = content_hash
    if source_url:
        entry["source_url"] = source_url

    index["documents"].append(entry)
    index["last_updated"] = datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ")
    _write_json(index_path, index)


# ═══════════════════════════════════════════
# Utility functions
# ═══════════════════════════════════════════

def _write_text(path: Path, content: str):
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


def _write_json(path: Path, data: dict):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _safe_filename(title: str, max_len: int = 40) -> str:
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '', title)
    safe = safe.strip().replace(" ", "-")
    return (safe[:max_len] if len(safe) > max_len else safe) or "untitled"


def _next_doc_id(docs_dir: Path) -> str:
    counter_path = docs_dir / ".counter"
    if counter_path.exists():
        try:
            counter = int(counter_path.read_text().strip())
        except ValueError:
            counter = 1
    else:
        counter = 1
    doc_id = f"DOC-{counter:03d}"
    counter_path.parent.mkdir(parents=True, exist_ok=True)
    counter_path.write_text(str(counter + 1))
    return doc_id


# ═══════════════════════════════════════════
# HTTP API（FastAPI）
# ═══════════════════════════════════════════

DEFAULT_DOCS_DIR = Path(os.environ.get(
    "LARKSCOUT_DOCS_DIR",
    os.path.expanduser("~/.larkscout/docs"),
))

MAX_UPLOAD_BYTES = int(os.environ.get("LARKSCOUT_MAX_UPLOAD_MB", "200")) * 1024 * 1024

_DOC_ID_RE = re.compile(r"^[A-Z]+-\d+$")
_TABLE_ID_RE = re.compile(r"^(table-)?\d+$")


def _validate_doc_id(doc_id: str) -> None:
    """Reject doc_id values that could cause path traversal."""
    if not _DOC_ID_RE.match(doc_id):
        raise HTTPException(400, f"invalid doc_id: {doc_id!r}")


def _validate_table_id(table_id: str) -> None:
    """Reject table_id values that could cause path traversal."""
    if not _TABLE_ID_RE.match(table_id):
        raise HTTPException(400, f"invalid table_id: {table_id!r}")


def _get_docs_dir() -> Path:
    d = DEFAULT_DOCS_DIR
    d.mkdir(parents=True, exist_ok=True)
    return d


# ---- Pydantic Models ----

class ParseResponse(BaseModel):
    doc_id: str
    filename: str
    file_type: str
    total_pages: int
    section_count: int
    table_count: int
    ocr_page_count: int
    digest: str
    manifest_path: str
    processing_time_sec: float


class SectionInfo(BaseModel):
    sid: str
    index: int
    title: str
    page_range: str
    char_count: int
    summary_preview: str = ""


class ManifestResponse(BaseModel):
    doc_id: str
    filename: str
    file_type: str | None = None
    source: str | None = None
    paths: dict[str, str]
    sections: list[dict[str, Any]]
    provenance: dict[str, Any] | None = None


class SearchResult(BaseModel):
    doc_id: str
    filename: str
    file_type: str
    digest: str
    tags: list[str] = []
    source: str = "upload"
    created_at: str | None = None
    score: float = 1.0


class SearchResponse(BaseModel):
    results: list[SearchResult]
    total: int


# ---- FastAPI app ----

app = FastAPI(title="Doc Reader API", version="3.0.0")


@app.get("/health")
async def health():
    return {
        "ok": True,
        "version": "3.0.0",
        "docs_dir": str(_get_docs_dir()),
        "supported_formats": ["pdf", "docx", "xlsx", "csv"],
    }


@app.post("/parse", response_model=ParseResponse)
async def api_parse_doc(
    file: UploadFile = File(...),
    doc_id: str | None = Form(None),
    generate_summary: bool = Form(True),
    force_ocr: bool = Form(False),
    ocr_pages: str | None = Form(None),
    extract_tables: bool = Form(True),
    max_tables_per_page: int = Form(3),
    concurrency: int = Form(3),
    tags: str | None = Form(None),  # JSON array string: '["Q3","financial"]'
    metadata: str | None = Form(None),  # JSON object string
):
    """Parse uploaded document (PDF/DOCX), return structured result."""
    docs_dir = _get_docs_dir()
    t0 = time.time()

    # Check format
    filename = file.filename or "unknown"
    suffix = Path(filename).suffix.lower()
    if suffix not in (".pdf", ".docx", ".xlsx", ".csv"):
        raise HTTPException(422, t("unsupported_format", fmt=suffix))

    # Parse tags
    parsed_tags: list[str] = []
    if tags:
        try:
            parsed_tags = json.loads(tags)
        except json.JSONDecodeError:
            parsed_tags = [t.strip() for t in tags.split(",") if t.strip()]

    # Validate doc_id if user-supplied
    if doc_id:
        _validate_doc_id(doc_id)

    # Save temp file
    d_id = doc_id or _next_doc_id(docs_dir)
    tmp_dir = docs_dir / d_id / ".tmp"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    tmp_path = tmp_dir / filename
    try:
        content = await file.read()
        if len(content) > MAX_UPLOAD_BYTES:
            raise HTTPException(413, f"file too large: {len(content)} bytes (max {MAX_UPLOAD_BYTES})")
        tmp_path.write_bytes(content)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, t("file_save_failed", err=str(e)))

    # Parse
    try:
        loop = asyncio.get_event_loop()
        if suffix == ".pdf":
            parsed = await loop.run_in_executor(None, lambda: parse_pdf(
                tmp_path, force_ocr=force_ocr, ocr_threshold=OCR_THRESHOLD,
                ocr_pages_spec=ocr_pages, extract_tables=extract_tables,
                max_tables_per_page=max_tables_per_page,
                concurrency=concurrency, cache_dir=docs_dir / d_id,
            ))
        elif suffix == ".docx":
            parsed = await loop.run_in_executor(None, lambda: parse_word(
                tmp_path, extract_tables=extract_tables,
            ))
        elif suffix == ".xlsx":
            parsed = await loop.run_in_executor(None, lambda: parse_xlsx(tmp_path))
        else:  # .csv
            parsed = await loop.run_in_executor(None, lambda: parse_csv(tmp_path))
    except Exception as e:
        raise HTTPException(500, t("parse_failed", err=str(e)))
    finally:
        # Cleanup temp file
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass

    # Summarize + write
    digest = t("summary_pending")
    try:
        if generate_summary:
            digest_text, brief_text, _ = await loop.run_in_executor(
                None, lambda: generate_summaries(parsed, concurrency=concurrency)
            )
            digest = digest_text
            await loop.run_in_executor(None, lambda: write_output(
                d_id, parsed, digest_text, brief_text, docs_dir,
                tags=parsed_tags, source="upload", original_path=str(filename),
            ))
        else:
            await loop.run_in_executor(None, lambda: write_output_extract_only(
                d_id, parsed, docs_dir, tags=parsed_tags, source="upload",
            ))
    except Exception as e:
        raise HTTPException(500, t("write_failed", err=str(e)))

    elapsed = round(time.time() - t0, 2)
    return ParseResponse(
        doc_id=d_id,
        filename=parsed.filename,
        file_type=parsed.file_type,
        total_pages=parsed.total_pages,
        section_count=len(parsed.sections),
        table_count=parsed.table_count,
        ocr_page_count=parsed.ocr_page_count,
        digest=digest[:300],
        manifest_path=f"docs/{d_id}/manifest.json",
        processing_time_sec=elapsed,
    )


# ---- Library query endpoints ----

@app.get("/library/search", response_model=SearchResponse)
async def library_search(
    q: str | None = None,
    tags: str | None = None,
    file_type: str | None = None,
    limit: int = 20,
):
    """Search document library."""
    docs_dir = _get_docs_dir()
    index_path = docs_dir / "doc-index.json"
    if not index_path.exists():
        return SearchResponse(results=[], total=0)

    try:
        index = json.loads(index_path.read_text(encoding="utf-8"))
    except Exception:
        return SearchResponse(results=[], total=0)

    documents = index.get("documents", [])

    # Filter
    if file_type:
        documents = [d for d in documents if d.get("file_type") == file_type]

    if tags:
        tag_list = [t.strip() for t in tags.split(",")]
        documents = [
            d for d in documents
            if any(t in (d.get("tags") or []) for t in tag_list)
        ]

    if q:
        q_lower = q.lower()
        scored = []
        for d in documents:
            score = 0.0
            if q_lower in (d.get("filename") or "").lower():
                score += 2.0
            if q_lower in (d.get("digest") or "").lower():
                score += 1.0
            for tag in d.get("tags") or []:
                if q_lower in tag.lower():
                    score += 1.5
            if score > 0:
                scored.append((d, score))
        scored.sort(key=lambda x: x[1], reverse=True)
        documents = [d for d, _ in scored[:limit]]
        scores = {d.get("id"): s for d, s in scored[:limit]}
    else:
        documents = documents[:limit]
        scores = {}

    results = [
        SearchResult(
            doc_id=d.get("id", ""),
            filename=d.get("filename", ""),
            file_type=d.get("file_type", ""),
            digest=d.get("digest", ""),
            tags=d.get("tags", []),
            source=d.get("source", "upload"),
            created_at=d.get("created_at"),
            score=scores.get(d.get("id"), 1.0),
        )
        for d in documents
    ]
    return SearchResponse(results=results, total=len(results))


@app.get("/library/{doc_id}/manifest")
async def get_manifest(doc_id: str):
    """Get document manifest."""
    _validate_doc_id(doc_id)
    p = _get_docs_dir() / doc_id / "manifest.json"
    if not p.exists():
        raise HTTPException(404, t("doc_not_found", doc_id=doc_id))
    return json.loads(p.read_text(encoding="utf-8"))


@app.get("/library/{doc_id}/digest")
async def get_digest(doc_id: str):
    """Get document digest (lowest token cost)."""
    _validate_doc_id(doc_id)
    p = _get_docs_dir() / doc_id / "digest.md"
    if not p.exists():
        raise HTTPException(404, t("digest_not_found", doc_id=doc_id))
    return {"doc_id": doc_id, "content": p.read_text(encoding="utf-8")}


@app.get("/library/{doc_id}/brief")
async def get_brief(doc_id: str):
    """Get document brief (medium token cost)."""
    _validate_doc_id(doc_id)
    p = _get_docs_dir() / doc_id / "brief.md"
    if not p.exists():
        raise HTTPException(404, t("brief_not_found", doc_id=doc_id))
    return {"doc_id": doc_id, "content": p.read_text(encoding="utf-8")}


@app.get("/library/{doc_id}/full")
async def get_full(doc_id: str):
    """Get full document text (high token cost, use sparingly)."""
    _validate_doc_id(doc_id)
    p = _get_docs_dir() / doc_id / "full.md"
    if not p.exists():
        raise HTTPException(404, t("full_not_found", doc_id=doc_id))
    return {"doc_id": doc_id, "content": p.read_text(encoding="utf-8")}


@app.get("/library/{doc_id}/section/{sid}")
async def get_section(doc_id: str, sid: str):
    """Read a single section by sid."""
    _validate_doc_id(doc_id)
    sections_dir = _get_docs_dir() / doc_id / "sections"
    if not sections_dir.exists():
        raise HTTPException(404, t("doc_not_found", doc_id=doc_id))

    # sid is in filename: 01-{sid}-{title}.md
    for f in sections_dir.iterdir():
        if f.is_file() and sid in f.name:
            return {"doc_id": doc_id, "sid": sid, "content": f.read_text(encoding="utf-8")}

    raise HTTPException(404, t("section_not_found", sid=sid))


@app.get("/library/{doc_id}/table/{table_id}")
async def get_table(doc_id: str, table_id: str):
    """Read a single table."""
    _validate_doc_id(doc_id)
    _validate_table_id(table_id)
    tables_dir = _get_docs_dir() / doc_id / "tables"
    if not tables_dir.exists():
        raise HTTPException(404, t("tables_dir_not_found", doc_id=doc_id))

    # table_id: "table-01" or "01"
    tid = table_id if table_id.startswith("table-") else f"table-{table_id}"
    p = tables_dir / f"{tid}.md"
    if not p.exists():
        raise HTTPException(404, t("table_not_found", table_id=table_id))
    return {"doc_id": doc_id, "table_id": table_id, "content": p.read_text(encoding="utf-8")}


@app.get("/library/{doc_id}/sections")
async def list_sections(doc_id: str):
    """List all sections from manifest."""
    _validate_doc_id(doc_id)
    manifest_path = _get_docs_dir() / doc_id / "manifest.json"
    if not manifest_path.exists():
        raise HTTPException(404, t("doc_not_found", doc_id=doc_id))
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    return {
        "doc_id": doc_id,
        "sections": manifest.get("sections", []),
    }


# ═══════════════════════════════════════════
# Startup
# ═══════════════════════════════════════════

if __name__ == "__main__":
    import uvicorn

    host = os.environ.get("HOST", "0.0.0.0")
    port = int(os.environ.get("PORT", "8090"))

    DEFAULT_DOCS_DIR.mkdir(parents=True, exist_ok=True)
    logger.info(f"LarkScout DocReader API v3.0 starting: {host}:{port}")
    logger.info(f"Docs directory: {DEFAULT_DOCS_DIR}")

    uvicorn.run(app, host=host, port=port)
